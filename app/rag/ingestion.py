import logging
from contextlib import suppress

from docx import Document
from langchain_core.documents import Document as LangchainDocument
from langchain_openai import OpenAIEmbeddings
from langchain_qdrant import QdrantVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams

from app.config import get_settings

logger = logging.getLogger(__name__)


def load_document(file_path: str) -> str:
    """Load text from DOCX file"""
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def ingest_document(
    document_path: str,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> None:
    """Ingest document into Qdrant vector database"""
    settings = get_settings()
    text_content = load_document(document_path)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=chunk_overlap
    )
    chunks = text_splitter.split_text(text_content)

    embeddings = OpenAIEmbeddings()
    client = QdrantClient(url=settings.QDRANT_URL)

    with suppress(Exception):
        client.create_collection(
            collection_name=settings.COLLECTION_NAME,
            vectors_config=VectorParams(size=1536, distance=Distance.COSINE),
        )

    vector_store = QdrantVectorStore(
        client=client, collection_name=settings.COLLECTION_NAME, embedding=embeddings
    )

    documents = [
        LangchainDocument(page_content=chunk, metadata={"source": document_path})
        for chunk in chunks
    ]
    vector_store.add_documents(documents)

    logger.info(f"Imported {len(chunks)} chunks from '{document_path}' to '{settings.COLLECTION_NAME}'")
